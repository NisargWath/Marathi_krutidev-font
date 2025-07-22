import streamlit as st
import speech_recognition as sr
import tempfile
import os
from indic_transliteration import sanscript
from indic_transliteration.sanscript import SchemeMap, SCHEMES, transliterate
from docx import Document
import base64
import io
from audio_recorder_streamlit import audio_recorder
import time
import cv2
import numpy as np
from PIL import Image
import pytesseract
import easyocr
import google.generativeai as genai
from datetime import datetime
import json

# Configure Tesseract path if needed (uncomment and modify path as per your system)
# For Windows:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# For Mac (if installed via Homebrew):
# pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'

# Hardcoded API key - Replace with your actual key
GEMINI_API_KEY = "AIzaSyC5yD31qrSfirPAyDZZd5UjPhZ4Nag93_0"  # Replace with your actual API key

# Unicode to Krutidev mapping dictionary (simplified version)
UNICODE_TO_KRUTIDEV = {
   "अ": "v", "आ": "vk", "इ": "b", "ई": "bZ", "उ": "m", "ऊ": "Å",
    "ऋ": "_", "ए": ",", "ऐ": ",s", "ओ": "vks", "औ": "vkS",
    "क": "d", "ख": "[k", "ग": "x", "घ": "?k", "ङ": "³",
    "च": "p", "छ": "N", "ज": "t", "झ": ">", "ञ": "¥",
    "ट": "V", "ठ": "B", "ड": "M", "ढ": "<", "ण": ".k",
    "त": "r", "थ": "Fk", "द": "n", "ध": "/k", "न": "u",
    "प": "i", "फ": "Q", "ब": "c", "भ": "Hk", "म": "e",
    "य": ";", "र": "j", "ल": "y", "व": "o", "श": "'k",
    "ष": "\"k", "स": "l", "ह": "g",
    "ा": "k", "ि": "f", "ी": "h", "ु": "q", "ू": "w", "े": "s",
    "ै": "S", "ो": "ks", "ौ": "kS", "ं": "M+", "ः": "%", "ँ": "~",
    "्": "",  # Halant - skip
    " ": " ", "\n": "\n"
}

class AIDocumentExtractor:
    def __init__(self):
        """Initialize Gemini Vision API"""
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            self.model = genai.GenerativeModel('gemini-1.5-flash')
            self.ready = True
        except Exception as e:
            st.warning(f"Gemini AI not available: {str(e)}")
            self.ready = False
    
    def extract_text_from_image(self, image, language="Hindi", document_type="Document"):
        """Extract text using Gemini Vision"""
        try:
            # Create extraction prompt based on language
            lang_instruction = "Hindi/Devanagari" if language == "Hindi" else "Marathi/Devanagari"
            
            prompt = f"""
            Extract ALL text from this {document_type} image in {lang_instruction} script.

            Please provide:
            1. Complete text extraction maintaining original structure
            2. Include all visible text: headers, body, numbers, dates
            3. Preserve formatting and line breaks where logical
            4. Extract both printed and handwritten text
            5. Maintain original {lang_instruction} characters
            6. If text is in Roman script, convert it appropriately

            Return only the extracted text without any additional commentary:
            """
            
            # Generate content using vision
            response = self.model.generate_content([prompt, image])
            
            if response and hasattr(response, 'text') and response.text:
                return response.text.strip()
            else:
                return "No text extracted from the image"
                
        except Exception as e:
            if "429" in str(e):
                return "Rate limit exceeded. Please wait and try again."
            elif "quota" in str(e).lower():
                return "API quota exceeded. Please check billing."
            else:
                return f"AI OCR Error: {str(e)}"

def convert_unicode_to_krutidev(unicode_text):
    """
    Convert Unicode Devanagari text to Krutidev font
    This is a basic implementation - for production use, consider using a specialized library
    """
    try:
        krutidev_text = ""
        i = 0
        while i < len(unicode_text):
            char = unicode_text[i]
            
            # Check for compound characters first
            if i < len(unicode_text) - 1:
                compound = char + unicode_text[i + 1]
                if compound in UNICODE_TO_KRUTIDEV:
                    krutidev_text += UNICODE_TO_KRUTIDEV[compound]
                    i += 2
                    continue
            
            # Single character conversion
            if char in UNICODE_TO_KRUTIDEV:
                krutidev_text += UNICODE_TO_KRUTIDEV[char]
            else:
                krutidev_text += char  # Keep non-Devanagari characters as-is
            
            i += 1
        
        return krutidev_text
    except Exception as e:
        # Fallback to transliteration if direct mapping fails
        try:
            return transliterate(unicode_text, sanscript.DEVANAGARI, sanscript.ITRANS)
        except:
            return f"Conversion error: {str(e)}"

# Function to preprocess image for better OCR results
def preprocess_image(image):
    """
    Preprocess the image to improve OCR accuracy
    """
    try:
        # Convert PIL image to OpenCV format
        open_cv_image = np.array(image)
        
        # Convert RGB to BGR if needed (OpenCV uses BGR)
        if len(open_cv_image.shape) == 3 and open_cv_image.shape[2] == 3:
            open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2BGR)
        
        # Convert to grayscale
        if len(open_cv_image.shape) == 3:
            gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
        else:
            gray = open_cv_image
        
        # Apply Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        
        # Apply threshold to get better contrast
        _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Morphological operations to clean up the image
        kernel = np.ones((1, 1), np.uint8)
        processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        
        return processed
    except Exception as e:
        st.error(f"Image preprocessing error: {str(e)}")
        return np.array(image)

# Function to extract text from image using Tesseract
def extract_text_tesseract(image, language='hin+eng'):
    """
    Extract text from image using Tesseract OCR
    language: 'hin+eng' for Hindi+English, 'mar+eng' for Marathi+English
    """
    try:
        # Preprocess the image
        processed_image = preprocess_image(image)
        
        # Configure Tesseract for Indian languages
        custom_config = f'--oem 3 --psm 6 -l {language}'
        
        # Extract text
        text = pytesseract.image_to_string(processed_image, config=custom_config)
        return text.strip()
    except Exception as e:
        return f"Tesseract OCR Error: {str(e)}"

# Function to extract text from image using EasyOCR
def extract_text_easyocr(image, languages=['hi', 'en']):
    """
    Extract text from image using EasyOCR
    languages: ['hi', 'en'] for Hindi+English, ['mr', 'en'] for Marathi+English
    """
    try:
        # Initialize EasyOCR reader
        reader = easyocr.Reader(languages, gpu=False)  # Set gpu=True if you have CUDA
        
        # Convert PIL image to numpy array
        image_array = np.array(image)
        
        # Extract text
        results = reader.readtext(image_array)
        
        # Combine all detected text
        extracted_text = ' '.join([result[1] for result in results])
        return extracted_text.strip()
    except Exception as e:
        return f"EasyOCR Error: {str(e)}"

# Function to recognize speech from audio file
def recognize_speech(audio_file, language='mr-IN'):
    """
    Recognize speech from audio file
    """
    r = sr.Recognizer()
    
    try:
        with sr.AudioFile(audio_file) as source:
            # Adjust for ambient noise
            r.adjust_for_ambient_noise(source)
            audio_data = r.record(source)
            
            # Use Google's speech recognition service
            text = r.recognize_google(audio_data, language=language)
            return text
    except sr.UnknownValueError:
        return "Could not understand audio"
    except sr.RequestError as e:
        return f"Error with the speech recognition service: {e}"
    except Exception as e:
        return f"Error: {str(e)}"

# Function to create a download link for a Word document
def create_download_link_docx(text, filename):
    """
    Create a download link for Word document
    """
    try:
        # Create a Word document
        doc = Document()
        
        # Add paragraph with Krutidev font
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'Kruti Dev 010'  # Set Krutidev font
        run.font.size = 12
        
        # Save document to a BytesIO object
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create a download link
        b64 = base64.b64encode(doc_io.read()).decode()
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📄 Download Word document</a>'
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return ""

def display_conversion_results(original_text, krutidev_text, source_type="text"):
    """
    Display conversion results with download options
    """
    # Show the original text
    st.subheader("📝 Original Text:")
    st.write(original_text)
    
    # Convert to Krutidev and display
    st.subheader("🔤 Text in Krutidev Format:")
    st.text_area("Krutidev Text", krutidev_text, height=150)
    
    # Create download options
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            label="📁 Download as Text File",
            data=krutidev_text,
            file_name=f"krutidev_{source_type}.txt",
            mime="text/plain"
        )
    
    with col2:
        # Create Word document download link
        word_link = create_download_link_docx(krutidev_text, f"krutidev_{source_type}.docx")
        if word_link:
            st.markdown(word_link, unsafe_allow_html=True)

def main():
    # Page configuration
    st.set_page_config(
        page_title="Enhanced Voice & Image to Krutidev Converter",
        page_icon="🔤",
        layout="wide"
    )
    
    st.title("🎙️📷🤖 Enhanced Voice and Image to Krutidev Font Converter")
    
    st.markdown("""
    ### Convert voice recordings or image text to Krutidev font with AI-powered OCR
    
    This enhanced application helps you convert:
    - **Voice recordings** in Hindi or Marathi to Krutidev text
    - **Audio files** to Krutidev text
    - **Images with text** to Krutidev text using three different OCR methods:
      - 🤖 **Google Gemini AI Vision** (Most accurate for complex documents)
      - 🔍 **EasyOCR** (Good for general text recognition)
      - 📝 **Tesseract OCR** (Traditional OCR method)
    
    Choose your preferred input method from the tabs below.
    """)
    
    # Initialize AI extractor
    ai_extractor = AIDocumentExtractor()
    
    # Language selection
    st.sidebar.header("⚙️ Settings")
    language = st.sidebar.radio(
        "Select Language:",
        ("Hindi", "Marathi"),
        help="Choose the language for voice recognition and OCR"
    )
    
    language_code = "hi-IN" if language == "Hindi" else "mr-IN"
    
    # AI OCR settings
    if ai_extractor.ready:
        st.sidebar.success("🤖 Gemini AI OCR: Ready")
    else:
        st.sidebar.warning("🤖 Gemini AI OCR: Not available")
        st.sidebar.caption("Update GEMINI_API_KEY in code to enable AI OCR")
    
    # Create tabs for different input methods
    tab1, tab2, tab3, tab4 = st.tabs([
        "🎙️ Record Audio", 
        "📁 Upload Audio", 
        "📷 Extract from Image",
        "✏️ Direct Text Input"
    ])
    
    # Tab 1: Record Audio
    with tab1:
        st.header("🎙️ Record Audio")
        st.write("Click the button below to record your voice directly in the browser:")
        
        # Using audio_recorder_streamlit component for recording
        audio_bytes = audio_recorder(
            text="Click to record",
            recording_color="#e8b62c",
            neutral_color="#6aa36f",
            icon_name="microphone-lines",
            icon_size="6x",
        )
        
        if audio_bytes:
            st.audio(audio_bytes, format="audio/wav")
            
            # Save the recorded audio to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_file:
                tmp_file.write(audio_bytes)
                recorded_file_path = tmp_file.name
            
            if st.button("🔍 Transcribe Recorded Audio", type="primary"):
                with st.spinner("🔄 Processing your recording..."):
                    # Recognize speech
                    transcribed_text = recognize_speech(recorded_file_path, language_code)
                    
                    if not transcribed_text.startswith("Error") and not transcribed_text.startswith("Could not"):
                        # Convert to Krutidev
                        krutidev_text = convert_unicode_to_krutidev(transcribed_text)
                        display_conversion_results(transcribed_text, krutidev_text, "recorded_audio")
                    else:
                        st.error(f"❌ {transcribed_text}")
                
                # Clean up the temporary file
                try:
                    os.unlink(recorded_file_path)
                except:
                    pass
    
    # Tab 2: Upload Audio
    with tab2:
        st.header("📁 Upload Audio File")
        st.write("Upload an audio file containing speech in Hindi or Marathi:")
        
        # File uploader for audio
        uploaded_file = st.file_uploader(
            "Choose an audio file", 
            type=["wav", "mp3", "m4a", "flac", "aac"],
            help="Supported formats: WAV, MP3, M4A, FLAC, AAC"
        )
        
        if uploaded_file is not None:
            # Display file info
            st.info(f"📄 File: {uploaded_file.name} ({uploaded_file.size} bytes)")
            
            # Save the uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_filepath = tmp_file.name
            
            st.audio(uploaded_file, format="audio/wav")
            
            if st.button("🔍 Transcribe Uploaded Audio", type="primary"):
                with st.spinner("🔄 Processing audio file..."):
                    # Recognize speech
                    transcribed_text = recognize_speech(tmp_filepath, language_code)
                    
                    if not transcribed_text.startswith("Error") and not transcribed_text.startswith("Could not"):
                        # Convert to Krutidev
                        krutidev_text = convert_unicode_to_krutidev(transcribed_text)
                        display_conversion_results(transcribed_text, krutidev_text, "uploaded_audio")
                    else:
                        st.error(f"❌ {transcribed_text}")
                
                # Clean up the temporary file
                try:
                    os.unlink(tmp_filepath)
                except:
                    pass
    
    # Tab 3: Extract from Image (Enhanced with AI OCR)
    with tab3:
        st.header("📷 Enhanced Text Extraction from Images")
        st.write("Upload an image containing text in Hindi or Marathi. Now with AI-powered OCR!")
        
        # File uploader for images
        uploaded_image = st.file_uploader(
            "Choose an image file", 
            type=["png", "jpg", "jpeg", "bmp", "tiff", "webp"],
            help="Supported formats: PNG, JPG, JPEG, BMP, TIFF, WebP"
        )
        
        if uploaded_image is not None:
            # Display the uploaded image
            image = Image.open(uploaded_image)
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.image(image, caption="Original Image", use_column_width=True)
            
            with col2:
                st.info(f"📄 File: {uploaded_image.name}")
                st.info(f"📐 Size: {image.size[0]} x {image.size[1]} pixels")
                
                # OCR method selection
                ocr_options = ["🤖 Gemini AI Vision (Recommended)", "🔍 EasyOCR", "📝 Tesseract OCR"]
                if not ai_extractor.ready:
                    ocr_options = ["🔍 EasyOCR (Recommended)", "📝 Tesseract OCR"]
                
                ocr_method = st.selectbox(
                    "Select OCR Method:",
                    ocr_options,
                    help="Gemini AI Vision is most accurate for complex documents and handwritten text"
                )
                
                # Document type for AI OCR
                if "Gemini AI" in ocr_method:
                    doc_type = st.selectbox(
                        "Document Type (for AI OCR):",
                        ["Document", "FIR", "Police Complaint", "Legal Document", "Handwritten Note", "Form"],
                        help="Helps AI better understand document structure"
                    )
            
            if st.button("🔍 Extract Text from Image", type="primary"):
                with st.spinner("🔄 Extracting text from image..."):
                    extracted_text = ""
                    
                    if "Gemini AI" in ocr_method and ai_extractor.ready:
                        # Use Gemini AI Vision
                        extracted_text = ai_extractor.extract_text_from_image(image, language, doc_type)
                        
                    elif "EasyOCR" in ocr_method:
                        # Set language codes for EasyOCR
                        if language == "Marathi":
                            ocr_languages = ['mr', 'en']  # Marathi and English
                        else:
                            ocr_languages = ['hi', 'en']  # Hindi and English
                        
                        extracted_text = extract_text_easyocr(image, ocr_languages)
                        
                    else:  # Tesseract OCR
                        # Set language codes for Tesseract
                        if language == "Marathi":
                            tesseract_lang = 'mar+eng'  # Marathi and English
                        else:
                            tesseract_lang = 'hin+eng'  # Hindi and English
                        
                        extracted_text = extract_text_tesseract(image, tesseract_lang)
                    
                    if extracted_text and not extracted_text.startswith("Error") and extracted_text.strip():
                        # Display extraction method used
                        method_name = ocr_method.split()[1] if "Gemini" in ocr_method else ocr_method.split()[0][2:]
                        st.success(f"✅ Text extracted using {method_name}")
                        
                        # Convert to Krutidev
                        krutidev_text = convert_unicode_to_krutidev(extracted_text)
                        display_conversion_results(extracted_text, krutidev_text, "image")
                        
                        # Show extraction stats
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("📊 Characters", len(extracted_text))
                        with col2:
                            st.metric("📝 Words", len(extracted_text.split()))
                        with col3:
                            st.metric("🔤 Method", method_name)
                            
                    else:
                        st.error("❌ Could not extract text from the image.")
                        st.write("**💡 Tips for better results:**")
                        st.write("- Use images with clear, high-contrast text")
                        st.write("- Ensure text is not too small or blurry")
                        st.write("- Try different OCR methods")
                        st.write("- For handwritten text, use Gemini AI Vision")
                        st.write("- Consider preprocessing the image for better contrast")
    
    # Tab 4: Direct Text Input
    with tab4:
        st.header("✏️ Direct Text Input")
        st.write("Enter or paste Unicode Devanagari text to convert to Krutidev:")
        
        # Text input area
        input_text = st.text_area(
            "Enter your text:",
            height=150,
            placeholder="Type or paste your Hindi/Marathi text here..."
        )
        
        if input_text.strip():
            if st.button("🔄 Convert to Krutidev", type="primary"):
                # Convert to Krutidev
                krutidev_text = convert_unicode_to_krutidev(input_text)
                display_conversion_results(input_text, krutidev_text, "direct_input")
    
    # Sidebar with additional information
    with st.sidebar:
        st.header("📋 Installation Guide")
        
        with st.expander("Required Packages"):
            st.code("""
pip install streamlit 
pip install speech-recognition 
pip install indic-transliteration 
pip install python-docx 
pip install audio-recorder-streamlit 
pip install opencv-python 
pip install pillow 
pip install pytesseract 
pip install easyocr
pip install google-generativeai
            """)
        
        with st.expander("Tesseract Setup"):
            st.write("""
            **Windows:**
            - Download from GitHub: UB-Mannheim/tesseract
            - Install Hindi/Marathi language packs
            
            **Linux:**
            ```bash
            sudo apt install tesseract-ocr
            sudo apt install tesseract-ocr-hin
            sudo apt install tesseract-ocr-mar
            ```
            
            **Mac:**
            ```bash
            brew install tesseract
            brew install tesseract-lang
            ```
            """)
        
        with st.expander("Gemini AI Setup"):
            st.write("""
            **Get API Key:**
            1. Visit Google AI Studio
            2. Create/login to Google account
            3. Generate API key
            4. Replace GEMINI_API_KEY in the code
            
            **Features:**
            - Most accurate for complex documents
            - Handles handwritten text well
            - Works with multiple languages
            - Best for legal documents, FIRs
            """)
        
        st.header("ℹ️ About")
        st.write("""
        Enhanced application converts text from various sources to Krutidev font format, 
        commonly used for Hindi and Marathi typing in India.
        
        **New Features:**
        - 🤖 AI-powered OCR with Google Gemini Vision
        - 📋 Document type detection
        - 🎯 Improved accuracy for handwritten text
        - 📊 Extraction statistics
        
        **All Features:**
        - Voice-to-text conversion
        - Audio file transcription  
        - Multi-method image text extraction
        - Direct text conversion
        - Multiple download formats
        """)
        
        st.header("🆘 Troubleshooting")
        with st.expander("Common Issues"):
            st.write("""
            **Audio not working:**
            - Check microphone permissions
            - Try uploading file instead
            
            **OCR not accurate:**
            - Use clearer, high-contrast images
            - Try Gemini AI for handwritten text
            - Try EasyOCR instead of Tesseract
            - Ensure proper language selection
            
            **AI OCR not working:**
            - Check API key configuration
            - Verify internet connection
            - Check API quota/billing
            
            **Conversion issues:**
            - Verify input text is in Unicode Devanagari
            - Check if special characters need manual adjustment
            """)

if __name__ == "__main__":
    main()